from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from lxml import etree
import copy
import os
import re


# ---------------------------------------------------------------------------
# Helpers de mesclagem de documentos externos
# ---------------------------------------------------------------------------

def _mesclar_docx_no_local(paragrafo_alvo, caminho_docx):
    """Substitui o parágrafo placeholder pelo conteúdo completo de um DOCX externo."""
    try:
        doc_fonte = Document(caminho_docx)
    except Exception as e:
        # Falha silenciosa: mantém o placeholder como texto de erro
        _substituir_texto_mantendo_formatacao(paragrafo_alvo, paragrafo_alvo.text,
                                              f"[Erro ao abrir DOCX: {e}]")
        return

    parent = paragrafo_alvo._element.getparent()
    if parent is None:
        return

    idx = list(parent).index(paragrafo_alvo._element)
    parent.remove(paragrafo_alvo._element)

    for child in list(doc_fonte.element.body):
        if child.tag.endswith('sectPr'):
            continue
        parent.insert(idx, copy.deepcopy(child))
        idx += 1


def _extrair_texto_pdf(caminho_pdf):
    """Extrai texto de um PDF usando pdfplumber (preferido) ou pypdf como fallback."""
    # Tentativa 1: pdfplumber (melhor qualidade de extração)
    try:
        import pdfplumber
        paginas = []
        with pdfplumber.open(caminho_pdf) as pdf:
            for page in pdf.pages:
                texto = page.extract_text()
                if texto and texto.strip():
                    paginas.append(texto.strip())
        return "\n\n".join(paginas)
    except ImportError:
        pass
    except Exception as e:
        return f"[Erro ao ler PDF com pdfplumber: {e}]"

    # Tentativa 2: pypdf
    try:
        from pypdf import PdfReader
        reader = PdfReader(caminho_pdf)
        paginas = []
        for page in reader.pages:
            texto = page.extract_text()
            if texto and texto.strip():
                paginas.append(texto.strip())
        return "\n\n".join(paginas)
    except ImportError:
        pass
    except Exception as e:
        return f"[Erro ao ler PDF com pypdf: {e}]"

    return f"[PDF: {os.path.basename(caminho_pdf)} — instale pdfplumber ou pypdf para extrair conteúdo]"


# ---------------------------------------------------------------------------
# Remoção da seção ARP
# ---------------------------------------------------------------------------

def _remover_secao_arp(doc):
    """
    Remove o Anexo ARP (Minuta da ATA de Registro de Preços) do documento.
    Deve ser chamado ANTES de _processar_paragrafo, pois o heading ARP tem
    texto verde que seria apagado por aquela função, impossibilitando a busca.
    """
    body = doc.element.body
    elementos_body = list(body)

    # Busca usando p.text da API python-docx (evita a triplicação de itertext())
    arp_element = None
    for p in doc.paragraphs:
        texto = p.text.strip().upper()
        if ("MINUTA DA ATA DE REGISTRO DE PREÇOS" in texto or
                "MINUTA DE ATA DE REGISTRO DE PREÇOS" in texto):
            if "...." not in texto and not re.search(r'\d+$', texto) and len(texto) < 200:
                arp_element = p._element
                break

    if arp_element is None:
        return

    # Encontra o ancestral direto do body (pode estar dentro de w:sdt)
    el = arp_element
    while el is not None and el.getparent() is not body:
        el = el.getparent()
    if el is None:
        return

    try:
        idx_corte = elementos_body.index(el)
    except ValueError:
        return

    # Sobe para encontrar o início da seção (page break ou "ANEXO")
    temp_idx = idx_corte
    limite_busca = max(0, idx_corte - 30)
    for j in range(idx_corte - 1, limite_busca - 1, -1):
        el_j = elementos_body[j]
        txt = re.sub(r'\s+', ' ', "".join(el_j.itertext()).strip().upper())
        try:
            xml_str = etree.tostring(el_j, encoding='unicode').upper()
        except Exception:
            xml_str = ""

        is_page_break = ('TYPE="PAGE"' in xml_str.replace(" ", "") or
                         "PAGEBREAKBEFORE" in xml_str.replace(" ", "") or
                         "BREAK" in xml_str)
        if is_page_break:
            temp_idx = j
            break
        if txt.startswith("ANEXO") and len(txt) < 80:
            temp_idx = j

    idx_corte = temp_idx

    # Remove elementos em branco imediatamente antes do corte
    while idx_corte > 0:
        el_prev = elementos_body[idx_corte - 1]
        txt = "".join(el_prev.itertext()).strip()
        try:
            xml_str = etree.tostring(el_prev, encoding='unicode').upper()
        except Exception:
            xml_str = ""

        is_page_break = ('TYPE="PAGE"' in xml_str.replace(" ", "") or
                         "PAGEBREAKBEFORE" in xml_str.replace(" ", "") or
                         "BREAK" in xml_str)
        if not txt or is_page_break:
            idx_corte -= 1
        else:
            break

    # Remove tudo a partir do corte (mantendo sectPr final)
    ultimo_elemento = elementos_body[-1] if elementos_body else None
    for el in elementos_body[idx_corte:]:
        if el is ultimo_elemento and el.tag.endswith('sectPr'):
            continue
        try:
            el.getparent().remove(el)
        except Exception:
            pass


# ---------------------------------------------------------------------------
# Preenchimento principal
# ---------------------------------------------------------------------------

def preencher_documento(caminho_modelo: str, caminho_saida: str, dados: dict) -> str:
    doc = Document(caminho_modelo)
    e_arp = dados.get("E_ARP", False)

    remover_amostra = dados.get("__REMOVER_AMOSTRA__", False)
    remover_vistoria = dados.get("__REMOVER_VISTORIA__", False)

    # ── Etapa 1: remover parágrafos de AMOSTRA / VISTORIA ─────────────────
    paragrafos_remover_set = set()
    for i, p in enumerate(doc.paragraphs):
        texto_upper = p.text.strip().upper()

        is_amostra = remover_amostra and "AMOSTRA" in texto_upper and len(texto_upper) < 60
        is_vistoria = remover_vistoria and "VISTORIA" in texto_upper and len(texto_upper) < 60

        if is_amostra or is_vistoria:
            paragrafos_remover_set.add(p)

            idx = i + 1
            while idx < len(doc.paragraphs):
                next_p = doc.paragraphs[idx]
                txt = next_p.text.strip()
                if not txt or (remover_amostra and "{{AMOSTRA}}" in txt) or (remover_vistoria and "{{VISTORIA}}" in txt):
                    paragrafos_remover_set.add(next_p)
                    idx += 1
                else:
                    break

            idx = i - 1
            while idx >= 0:
                prev_p = doc.paragraphs[idx]
                if not prev_p.text.strip():
                    paragrafos_remover_set.add(prev_p)
                    idx -= 1
                else:
                    break

    inseriu_linha_embranco = False
    for p in paragrafos_remover_set:
        try:
            parent = p._element.getparent()
            if parent is not None:
                idx = list(parent).index(p._element)
                parent.remove(p._element)
                if not inseriu_linha_embranco:
                    novo_p = OxmlElement("w:p")
                    parent.insert(idx, novo_p)
                    inseriu_linha_embranco = True
        except Exception:
            pass

    # ── Etapa 2: mesclar documentos externos (DFD, ETP, TR) ───────────────
    # Deve ocorrer ANTES de _processar_paragrafo para que os placeholders
    # ainda existam no texto.
    PLACEHOLDERS_DOC = {"{{DFD}}", "{{ETP}}", "{{TR}}"}
    chaves_doc_processadas = set()

    for chave_ph in PLACEHOLDERS_DOC:
        caminho_arq = str(dados.get(chave_ph, "")).strip()
        if not caminho_arq or not os.path.exists(caminho_arq):
            continue  # será tratado como substituição de texto vazio depois

        ext = os.path.splitext(caminho_arq)[1].lower()
        chaves_doc_processadas.add(chave_ph)

        for p in list(doc.paragraphs):
            if chave_ph in p.text:
                if ext == '.docx':
                    _mesclar_docx_no_local(p, caminho_arq)
                elif ext == '.pdf':
                    texto_pdf = _extrair_texto_pdf(caminho_arq)
                    _inserir_multilinhas_safe(p, chave_ph, texto_pdf)
                else:
                    # Tipo não suportado: coloca só o nome do arquivo
                    _substituir_texto_mantendo_formatacao(p, chave_ph, os.path.basename(caminho_arq))
                break  # cada placeholder aparece uma única vez

        # Zera no dict para que _processar_paragrafo não tente substituir
        # o placeholder que já foi consumido
        dados = dict(dados)
        dados[chave_ph] = ""

    # ── Etapa 3: remover seção ARP ANTES de processar parágrafos ──────────
    # CRÍTICO: _processar_paragrafo limpa o texto verde do heading ARP,
    # o que tornaria impossível localizar a seção depois. Por isso a remoção
    # deve acontecer aqui, enquanto o texto ainda está intacto.
    if not e_arp:
        _remover_secao_arp(doc)

    # ── Etapa 4: substituir placeholders e remover texto verde ─────────────
    for paragrafo in list(doc.paragraphs):
        _processar_paragrafo(paragrafo, dados, e_arp)

    for tabela in doc.tables:
        colunas_para_remover = []
        for i, coluna in enumerate(tabela.columns):
            remover_esta = False
            for cell in coluna.cells:
                if "__REMOVER_COLUNA__" in cell.text:
                    remover_esta = True
                    break
            if remover_esta:
                colunas_para_remover.append(i)

        for col_idx in reversed(colunas_para_remover):
            for row in tabela.rows:
                try:
                    tc = row.cells[col_idx]._tc
                    tc.getparent().remove(tc)
                except Exception:
                    pass

        for linha in tabela.rows:
            for celula in linha.cells:
                for paragrafo in list(celula.paragraphs):
                    _processar_paragrafo(paragrafo, dados, e_arp)

    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header:
                for paragrafo in list(header.paragraphs):
                    _processar_paragrafo(paragrafo, dados, e_arp)
                for tabela in header.tables:
                    for linha in tabela.rows:
                        for celula in linha.cells:
                            for paragrafo in list(celula.paragraphs):
                                _processar_paragrafo(paragrafo, dados, e_arp)
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer:
                for paragrafo in list(footer.paragraphs):
                    _processar_paragrafo(paragrafo, dados, e_arp)
                for tabela in footer.tables:
                    for linha in tabela.rows:
                        for celula in linha.cells:
                            for paragrafo in list(celula.paragraphs):
                                _processar_paragrafo(paragrafo, dados, e_arp)

    # ── Etapa 5: limpar realce verde residual ──────────────────────────────
    def limpar_realce_verde(elemento_raiz):
        for node in elemento_raiz.xpath('.//w:highlight'):
            val = node.get(qn('w:val'))
            if val in ['green', 'brightGreen']:
                parent = node.getparent()
                if parent is not None:
                    parent.remove(node)

    limpar_realce_verde(doc.element)
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header:
                limpar_realce_verde(header._element)
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer:
                limpar_realce_verde(footer._element)

    doc.save(caminho_saida)
    return caminho_saida


# ---------------------------------------------------------------------------
# Substituição de texto mantendo formatação
# ---------------------------------------------------------------------------

def _substituir_texto_mantendo_formatacao(paragrafo, marcador, valor_str):
    if marcador not in paragrafo.text:
        return False

    texto_p = paragrafo.text
    idx_inicio = texto_p.find(marcador)

    while idx_inicio != -1:
        idx_fim = idx_inicio + len(marcador)
        pos_atual = 0
        runs_envolvidos = []

        for run in paragrafo.runs:
            texto_run = run.text
            len_run = len(texto_run)
            pos_fim_run = pos_atual + len_run

            if pos_fim_run > idx_inicio and pos_atual < idx_fim:
                runs_envolvidos.append({
                    'run': run,
                    'inicio_run': pos_atual,
                    'fim_run': pos_fim_run,
                    'texto_original': texto_run
                })

            pos_atual = pos_fim_run

        if runs_envolvidos:
            for j, info in enumerate(runs_envolvidos):
                run = info['run']
                txt = info['texto_original']

                inicio_intersecao = max(0, idx_inicio - info['inicio_run'])
                fim_intersecao = min(len(txt), idx_fim - info['inicio_run'])

                prefixo = txt[:inicio_intersecao]
                sufixo = txt[fim_intersecao:]

                if j == 0:
                    run.text = prefixo + valor_str + sufixo
                else:
                    run.text = prefixo + sufixo

        texto_p = paragrafo.text
        idx_inicio = texto_p.find(marcador)

    return True


# ---------------------------------------------------------------------------
# Formatação markdown (bold / bold+italic)
# ---------------------------------------------------------------------------

def _aplicar_formatacao_markdown_avancado(paragrafo):
    if "***" not in paragrafo.text and "**" not in paragrafo.text:
        return

    runs = list(paragrafo.runs)
    for run in runs:
        texto_run = run.text
        if "***" in texto_run or "**" in texto_run:
            parts = re.split(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*)', texto_run, flags=re.DOTALL)
            if len(parts) > 1:
                run.text = ""
                run_base_element = run._element

                for part in parts:
                    if not part:
                        continue

                    new_run = paragrafo.add_run()
                    new_run._element.getparent().remove(new_run._element)
                    run_base_element.addnext(new_run._element)
                    run_base_element = new_run._element

                    if run._element.rPr is not None:
                        new_run._element.insert(0, copy.deepcopy(run._element.rPr))

                    if part.startswith('***') and part.endswith('***'):
                        new_run.text = part[3:-3]
                        new_run.font.bold = True
                        new_run.font.italic = True
                        new_run.font.underline = True
                    elif part.startswith('**') and part.endswith('**'):
                        new_run.text = part[2:-2]
                        new_run.font.bold = True
                    else:
                        new_run.text = part

                try:
                    run._element.getparent().remove(run._element)
                except Exception:
                    pass


# ---------------------------------------------------------------------------
# Inserção de texto multilinhas (respeita \n e linhas em branco)
# ---------------------------------------------------------------------------

def _inserir_multilinhas_safe(paragrafo, marcador, valor_str):
    texto_p = paragrafo.text
    idx_inicio = texto_p.find(marcador)
    rPr_base = None

    if idx_inicio != -1:
        pos_atual = 0
        for run in paragrafo.runs:
            len_run = len(run.text)
            if pos_atual + len_run > idx_inicio:
                if run._element.rPr is not None:
                    rPr_base = copy.deepcopy(run._element.rPr)
                break
            pos_atual += len_run

    if rPr_base is None and paragrafo.runs:
        for r in paragrafo.runs:
            if r._element.rPr is not None:
                rPr_base = copy.deepcopy(r._element.rPr)
                break

    valor_str = str(valor_str).replace('\r', '')
    # ── FIX Bug 1: só colapsa 3+ newlines consecutivos (preserva linhas em branco
    #    simples que geram parágrafos vazios — essenciais nos blocos de assinatura)
    valor_str = re.sub(r'\n{3,}', '\n\n', valor_str)
    # ── FIX Bug 1: NÃO filtra linhas vazias — elas viram parágrafos em branco no Word
    linhas = valor_str.strip('\n').split('\n')

    if not linhas or (len(linhas) == 1 and linhas[0] == ""):
        _substituir_texto_mantendo_formatacao(paragrafo, marcador, "")
        return []

    _substituir_texto_mantendo_formatacao(paragrafo, marcador, linhas[0])

    paragrafo_atual = paragrafo
    novos_paragrafos = []

    for linha in linhas[1:]:
        novo_p = OxmlElement("w:p")
        paragrafo_atual._p.addnext(novo_p)

        if paragrafo_atual._p.pPr is not None:
            novo_p.append(copy.deepcopy(paragrafo_atual._p.pPr))

        novo_paragrafo_obj = Paragraph(novo_p, paragrafo_atual._parent)
        novo_run = novo_paragrafo_obj.add_run(linha)

        if rPr_base is not None:
            novo_run._element.insert(0, copy.deepcopy(rPr_base))

        paragrafo_atual = novo_paragrafo_obj
        novos_paragrafos.append(novo_paragrafo_obj)

    return novos_paragrafos


# ---------------------------------------------------------------------------
# Processamento de um parágrafo (substituição + remoção de verde)
# ---------------------------------------------------------------------------

def _processar_paragrafo(paragrafo, dados, e_arp):
    apagou_algo = False

    for run in paragrafo.runs:
        is_green = False
        highlight_element = None
        if run._element.rPr is not None:
            highlight = run._element.rPr.find(qn('w:highlight'))
            if highlight is not None:
                val = highlight.get(qn('w:val'))
                if val in ['green', 'brightGreen']:
                    is_green = True
                    highlight_element = highlight

        if is_green:
            if not e_arp:
                if run.text:
                    run.text = ""
                    apagou_algo = True
            else:
                if highlight_element is not None:
                    run._element.rPr.remove(highlight_element)

    paragrafos_para_processar = [paragrafo]

    for chave, valor in dados.items():
        if chave in ["E_ARP", "__REMOVER_AMOSTRA__", "__REMOVER_VISTORIA__"]:
            continue

        marcador = chave if chave.startswith("{{") and chave.endswith("}}") else f"{{{{{chave}}}}}"

        if isinstance(valor, list):
            valor_str = "\n".join([str(v) for v in valor if str(v).strip()])
        else:
            valor_str = str(valor)

        novos_adicionados = []
        for p_atual in list(paragrafos_para_processar):
            if marcador in p_atual.text:
                if '\n' in valor_str:
                    novos_ps = _inserir_multilinhas_safe(p_atual, marcador, valor_str)
                    novos_adicionados.extend(novos_ps)
                    if not valor_str.strip() and not p_atual.text.strip():
                        apagou_algo = True
                else:
                    _substituir_texto_mantendo_formatacao(p_atual, marcador, valor_str)
                    if not valor_str.strip() and not p_atual.text.strip():
                        apagou_algo = True

        paragrafos_para_processar.extend(novos_adicionados)

    for p_atual in paragrafos_para_processar:
        _aplicar_formatacao_markdown_avancado(p_atual)

        if apagou_algo and not p_atual.text.strip():
            try:
                p_element = p_atual._element
                p_element.getparent().remove(p_element)
            except Exception:
                pass